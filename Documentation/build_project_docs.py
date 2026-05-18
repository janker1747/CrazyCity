from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\kin4e\Desktop\CrazyCityrep\CrazyCity\Documentation\CrazyCity_Project_Documentation.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E79")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in tbl.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return tbl


def note(doc, title, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "EAF2F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_code_ref(doc, path):
    p = doc.add_paragraph()
    p.style = "Intense Quote"
    r = p.add_run(path)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def h(doc, level, text):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Calibri"
    styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 1"].font.size = Pt(18)
styles["Heading 2"].font.size = Pt(14)
styles["Heading 3"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CrazyCity\nТехническая документация проекта")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Архитектура модулей, связи систем и новые правила Cargo System")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
note(
    doc,
    "Назначение документа",
    "Это рабочее руководство по текущему состоянию проекта. Оно объясняет, какие модули существуют, "
    "как они связаны, где настраиваются данные, и как безопасно расширять новые грузовые механики.",
)

h(doc, 1, "1. Карта проекта")
doc.add_paragraph(
    "Основной код находится в Assets/2_script. Проект построен вокруг игрока-машины, системы очков, "
    "грузов, бустов, врагов-полицейских, UI, выбора машины и трюков."
)
table(
    doc,
    ["Модуль", "Папка / файлы", "Ответственность", "Связан с"],
    [
        ["Player", "Assets/2_script/Player", "Главный runtime-объект игрока, ввод, очки, бусты, груз, физические множители.", "Cargo, Boosts, UI, Enemy, Tricks"],
        ["Cargo System", "Assets/2_script/Cargo", "Подбор нескольких грузов, таймеры, доставка, комбо, эффекты грузов, стрелка, UI грузов.", "Player, ScoreSystem, Enemy, DOTween"],
        ["Boosts", "Assets/2_script/BoostsLogic", "Подбираемые бонусы: щит, скорость, x2 очки, time stop, power collision, stone wand.", "Player, ScoreSystem, TimerUI"],
        ["Enemy", "Assets/2_script/Enemy]", "AI преследования и штрафы от столкновения с полицейскими.", "Player, Cargo, ArcadeVehicleController"],
        ["Score/UI", "Assets/2_script/Player/PlayerUI", "Очки, таймеры, иконки бустов, отображение скорости.", "Player, Boosts, Cargo UI"],
        ["Car Loader", "Assets/2_script/CarLoader", "Выбор машины в гараже и спавн выбранного префаба игрока.", "Scenes, UI"],
        ["Tricks", "Assets/2_script/Trick", "Трюки в воздухе, начисление очков, футбольный проп.", "PlayerAirController, ScoreSystem"],
        ["Impact/Particles", "Assets/2_script/ParticleAndOtherVisual", "Реакции на столкновения, частицы, звук, начисление/штраф очков.", "PlayerCollisionHandler"],
        ["Environment", "Assets/2_script/Enviiroment", "Зоны и jump pad.", "Player/Rigidbody"],
    ],
    [3.0, 4.2, 6.2, 4.2],
)

h(doc, 1, "2. Runtime-поток игры")
doc.add_paragraph("Упрощенная цепочка выполнения в основной сцене выглядит так:")
numbered(doc, "CarSelectionManager хранит выбранную машину между сценами.")
numbered(doc, "CarGameLoader спавнит Player prefab в одной из точек спавна.")
numbered(doc, "Player в Awake создает BoostSlot, BoostSystem, ScoreSystem и инициализирует PlayerCargoModule.")
numbered(doc, "Каждый Update: BoostSystem обновляет активные timed-бусты, PlayerCargoModule обновляет все активные грузы.")
numbered(doc, "CargoPickup добавляет груз игроку, DeliveryPoint сдает все активные грузы на общей точке доставки.")
numbered(doc, "ScoreSystem рассылает события UI и camera feedback при изменении очков.")

h(doc, 1, "3. Player")
add_code_ref(doc, "Assets/2_script/Player/Player.cs")
doc.add_paragraph(
    "Player является центральным фасадом для большинства систем. Он держит ссылки на физику машины, UI, очки, бусты и Cargo System. "
    "После последних изменений Player также хранит стакающиеся множители скорости, ускорения, гравитации и downforce."
)
table(
    doc,
    ["Часть", "Что делает", "Кто использует"],
    [
        ["BoostSlot / BoostSystem", "Хранит один подобранный буст и обновляет активные эффекты.", "BoostPickup, UiPlayer"],
        ["ScoreSystem", "Добавляет/отнимает очки, применяет общий multiplier.", "Cargo, Tricks, Impact, Enemy"],
        ["CargoModule", "Принимает грузы, тикает таймеры, сдает/проваливает доставки.", "CargoPickup, DeliveryPoint, CargoInventoryUI"],
        ["Vehicle multipliers", "Позволяют грузам менять MaxSpeed, accelaration, gravity, downforce без перезаписи друг друга.", "HeavyCargo, AntiGravityCargo, MoonStoneCargo"],
        ["Cargo collision hooks", "Пробрасывает столкновения и score damage грузам.", "PlayerCollisionHandler, EnemyCollisionHandler"],
    ],
    [3.3, 7.0, 5.0],
)
note(
    doc,
    "Важно",
    "Старые методы AddSpeed/EndBonusSpeed сохранены для SpeedBoost, но теперь они работают через общий пересчет vehicle-множителей. "
    "Если новый груз меняет характеристики машины, используйте Add/Remove методы множителей, а не прямую запись в ArcadeVehicleController.",
)

h(doc, 1, "4. Cargo System после переработки")
add_code_ref(doc, "Assets/2_script/Cargo")
doc.add_paragraph(
    "Система больше не завязана на один currentCargo. Игрок может везти любое количество грузов. "
    "Точка доставки остается одной общей: она создается при первом активном грузе и удаляется, когда грузов больше нет."
)
table(
    doc,
    ["Класс", "Роль", "Ключевые связи"],
    [
        ["Cargo", "Базовый ScriptableObject груза. Хранит имя, иконку, базовую стоимость, deliveryTime, comboAmount и виртуальные хуки.", "Наследники грузов"],
        ["ActiveCargo", "Runtime-контейнер конкретного активного груза: Cargo, ElapsedTime, DamageMultiplier, таймеры состояния.", "PlayerCargoModule"],
        ["PlayerCargoModule", "Список активных грузов, tick каждого груза, сдача, провал, комбо, события UI.", "Player, CargoPickup, DeliveryPoint, CargoInventoryUI"],
        ["CargoPickup", "Trigger-пикап. Регистрируется в статическом списке ActivePickups для стрелки.", "Player.TryTakeCargo, CargoArrowUI"],
        ["CargoManager", "Создает и удаляет одну общую DeliveryPoint.", "PlayerCargoModule"],
        ["DeliveryPoint", "При входе игрока сдает все активные грузы.", "Player.CompleteDelivery(true)"],
        ["CargoArrowUI", "Ищет ближайший активный CargoPickup и поворачивает стрелку к нему.", "CargoPickup.ActivePickups"],
        ["CargoInventoryUI", "Показывает иконки активных грузов и счетчик комбо через DOTween.", "PlayerCargoModule events"],
    ],
    [3.4, 7.2, 4.7],
)

h(doc, 2, "4.1 Жизненный цикл груза")
numbered(doc, "Игрок входит в CargoPickup.")
numbered(doc, "CargoPickup вызывает Player.TryTakeCargo(cargoData).")
numbered(doc, "PlayerCargoModule создает ActiveCargo и добавляет его в activeCargos.")
numbered(doc, "Cargo.OnPickup(player, module, activeCargo) запускает индивидуальную механику груза.")
numbered(doc, "Каждый Tick груз получает независимый ElapsedTime и вызов Cargo.OnTick.")
numbered(doc, "Если DeliveryTime истек, проваливается только этот груз.")
numbered(doc, "Если игрок заезжает в DeliveryPoint, сдаются все активные грузы пачкой.")
numbered(doc, "Награда считается как сумма value всех грузов * сумма ComboAmount * глобальные reward-множители.")

h(doc, 2, "4.2 Расширяемые хуки Cargo")
table(
    doc,
    ["Хук", "Когда вызывается", "Для чего использовать"],
    [
        ["OnPickup(Player)", "При подборе. Старый API сохранен.", "Простые эффекты без доступа к ActiveCargo."],
        ["OnPickup(Player, PlayerCargoModule, ActiveCargo)", "При подборе нового груза.", "Сложные эффекты, которым нужен runtime state или модуль."],
        ["OnTick(..., deltaTime)", "Каждый кадр для каждого активного груза.", "Периодический урон, AddForce, случайные события, деградация стоимости."],
        ["OnPlayerCollision(..., Collision)", "Когда PlayerCollisionHandler ловит столкновение.", "Хрупкость, электричество, реакция на удар."],
        ["OnPlayerScoreDamage(..., damage)", "Когда враг снял очки с игрока.", "VIP, уничтожение при уроне."],
        ["GetTimerScaleForOtherCargo", "При расчете скорости таймера другого груза.", "Стабилизатор и нервный груз."],
        ["GetGlobalRewardMultiplier", "Перед финальной выдачей награды за пачку.", "Золотой груз."],
        ["ModifyScoreDamage", "Перед снятием очков врагом.", "Бронированный груз."],
        ["ProvidesCargoProtection", "Когда другой груз проверяет защиту.", "VIP, броня, щитовые механики."],
        ["CalculateValue(Player, module, activeCargo)", "При сдаче.", "Награды, зависящие от других активных грузов."],
    ],
    [4.2, 4.0, 7.0],
)

h(doc, 2, "4.3 Добавленные типы грузов")
table(
    doc,
    ["Груз", "Поведение", "Настройки"],
    [
        ["FragileCargo", "Теряет DamageMultiplier при сильном столкновении, может провалиться.", "damageVelocity, valueLossPerHit"],
        ["HeavyCargo", "Уменьшает скорость и ускорение, эффект стакается и снимается при сдаче/провале.", "speedMultiplier, accelerationMultiplier"],
        ["RadioactiveCargo", "Периодически отнимает очки, но имеет повышенную стоимость.", "scoreDrainInterval, scoreDrain, payoutMultiplier"],
        ["ContrabandCargo", "Повышенная награда и спавн копов при подборе.", "policePrefab, spawnRadius, copsToSpawn"],
        ["ElectricCargo", "Периодически замораживает ближайшего копа; замораживает копа при столкновении.", "shockRadius, shockInterval, freezeDuration"],
        ["LivingCargo", "В случайный момент подбрасывает игрока.", "minThrowDelay, maxThrowDelay, throwForce"],
        ["VipCargo", "Проваливается при получении score damage; дает cargo-защиту другим грузам.", "-"],
        ["CriticalCargo", "Теряет сохранность, если скорость игрока ниже 10.", "requiredSpeed, valueLossPerSecond"],
        ["AntiGravityCargo", "Уменьшает gravity и downforce машины.", "gravityMultiplier, downforceMultiplier"],
        ["StabilizerCargo", "Замедляет таймеры других грузов.", "otherCargoTimerMultiplier"],
        ["NervousCargo", "Ускоряет таймеры других грузов.", "otherCargoTimerMultiplier"],
        ["PairedCargo", "Дает награду только если активен второй такой же cargo asset.", "pairedRewardMultiplier"],
        ["CannedBoxCargo", "Дает +4 в comboAmount.", "ComboAmount = 4"],
        ["GoldCargo", "Увеличивает награду за все активные доставки.", "allCargoRewardMultiplier"],
        ["ArmoredCargo", "Уменьшает score damage от врагов и считается защитой.", "damageMultiplier"],
        ["FragileElectronicsCargo", "Ломается даже от слабого столкновения.", "breakVelocity"],
        ["CrystalBoxCargo", "Огромная награда; без защиты уничтожает все активные грузы.", "payoutMultiplier"],
        ["RocketCargo", "Периодически толкает игрока вперед AddForce.", "forceInterval, force"],
    ],
    [3.7, 7.5, 4.0],
)
note(
    doc,
    "Настройка новых cargo assets",
    "Создавайте asset через Create > Cargo > нужный тип. После этого назначайте asset в CargoPickup.cargoData. "
    "Для ContrabandCargo обязательно назначьте policePrefab, иначе копы не появятся.",
)

h(doc, 1, "5. UI")
doc.add_paragraph("UI разделен на несколько простых компонентов.")
table(
    doc,
    ["Класс", "Файл", "Назначение"],
    [
        ["UiPlayer", "Player/PlayerUI/UiPlayer.cs", "Скорость, кнопка использования буста, иконки активных бустов."],
        ["ScoreUI", "Player/PlayerUI/ScoreUI.cs", "Подписывается на ScoreSystem.OnScoreChanged и обновляет текст очков."],
        ["TimerUI", "Player/PlayerUI/TimerUI.cs", "Круговой таймер для timed-бустов, анимации через DOTween."],
        ["CargoInventoryUI", "Cargo/CargoInventoryUI.cs", "Новая панель активных грузов: иконки грузов и счетчик комбо."],
        ["CargoArrowUI", "Cargo/CargoArrowUI.cs", "Стрелка на ближайший доступный груз."],
        ["GarageUIManager", "CarLoader/GarageUIManager.cs", "UI выбора машины, слайдеры характеристик, DOTween-анимации."],
    ],
    [3.4, 5.0, 7.2],
)
h(doc, 2, "5.1 Подключение CargoInventoryUI")
numbered(doc, "Создайте объект UI-панели в Canvas.")
numbered(doc, "Добавьте CargoInventoryUI.")
numbered(doc, "Назначьте cargoModule вручную или оставьте пустым: компонент попробует найти Player.")
numbered(doc, "Назначьте iconsRoot - контейнер для иконок.")
numbered(doc, "Назначьте cargoIconPrefab - Image prefab для одной иконки.")
numbered(doc, "Назначьте comboText - TMP_Text для текущего comboAmount.")

h(doc, 1, "6. Boost System")
doc.add_paragraph(
    "Бусты построены через интерфейсы IBoost, ITimedBoost и IEventBoost. BoostData является ScriptableObject-фабрикой, "
    "создающей runtime-объект буста для конкретного Player."
)
table(
    doc,
    ["Компонент", "Роль"],
    [
        ["BoostPickup", "При trigger enter кладет BoostData в PlayerBoostSlot."],
        ["PlayerBoostSlot", "Хранит один текущий буст и отправляет событие BoostPickUP со sprite для UI."],
        ["BoostSystem", "Активирует буст, тикает ITimedBoost, подписывает IEventBoost."],
        ["DoublePointBoost", "Меняет ScoreSystem multiplier на 2 и возвращает 1 после таймера."],
        ["SpeedBoost", "Добавляет временный бонус скорости через Player.AddSpeed."],
        ["ShieldBoost", "Включает щит игрока до первого вражеского удара."],
        ["SafeBoost", "Фиксирует минимальные очки, ниже которых score не упадет."],
        ["PowerCollision", "Разрешает силовое столкновение с Police через PlayerCollisionHandler."],
        ["StoneWandBoost", "Спавнит камень, который бьет IHittable/ImpactSource."],
        ["TimeStopBoost", "Работает через TimeStopManager и rigidbody целей."],
    ],
    [4.0, 10.8],
)

h(doc, 1, "7. Enemy и столкновения")
doc.add_paragraph(
    "Enemy-система сейчас минимальная. AICarChase управляет машиной через ArcadeVehicleController и NavMeshAgent. "
    "EnemyCollisionHandler снимает очки с игрока, но теперь перед этим пропускает урон через cargo-модификаторы."
)
table(
    doc,
    ["Класс", "Что делает", "Важные связи"],
    [
        ["AICarChase", "Строит путь к player target и передает override input в ArcadeVehicleController.", "NavMeshAgent, ArcadeVehicleController"],
        ["EnemyCollisionHandler", "При столкновении с Player проверяет shield, затем снимает score damage.", "Player.ModifyCargoScoreDamage, NotifyCargoScoreDamage"],
        ["CargoFreezeEffect", "Временный helper для ElectricCargo: делает rigidbody isKinematic, затем восстанавливает.", "ElectricCargo"],
    ],
    [3.8, 6.0, 5.5],
)

h(doc, 1, "8. Очки, столкновения и эффекты")
doc.add_paragraph(
    "ScoreSystem - чистый C# класс без MonoBehaviour. Он хранит текущие очки, multiplier и режим safe. "
    "Изменения очков рассылаются через события, на которые подписаны UI и камера."
)
table(
    doc,
    ["Поток", "Описание"],
    [
        ["AddScore", "Умножает amount на текущий multiplier и вызывает OnScoreChanged/OnScoreAdded."],
        ["MinusScore", "Отнимает amount, учитывает SafeBoost, вызывает OnScoreChanged/OnScoreRemoved."],
        ["ImpactSystem", "На OnImpact проигрывает частицы/звук, добавляет или отнимает score из ImpactData."],
        ["CameraScoreFeedback", "Слушает OnScoreAdded/OnScoreRemoved и двигает Cinemachine offset через DOTween."],
    ],
    [4.0, 10.8],
)

h(doc, 1, "9. Выбор машины и гараж")
doc.add_paragraph(
    "Данные машины лежат в CarItemSO: prefab игрока, sprite, имя и числовые характеристики для UI. "
    "CarSelectionManager является singleton и не уничтожается при смене сцены."
)
note(
    doc,
    "Проверить в коде",
    "В CarGameLoader.Awake сейчас стоит `if (_dontTryLoad = true)`, то есть присваивание вместо сравнения. "
    "Такой код всегда ставит _dontTryLoad в true и выходит из Awake. Если спавн игрока не работает, это первое место для исправления.",
)

h(doc, 1, "10. Tricks и пропы")
doc.add_paragraph(
    "Трюки активируются только в воздухе. PlayerAirController слушает ArcadeVehicleController.OnGrounded, "
    "а затем по Q/E/Space пытается стартовать TrickData из PlayerTrickLoadout."
)
table(
    doc,
    ["Класс", "Роль"],
    [
        ["TrickData", "ScriptableObject с названием, animatorTrigger, иконкой, score и duration."],
        ["PlayerTrickLoadout", "Связывает клавиши с TrickData."],
        ["PlayerAirController", "Стартует анимацию, завершает или проваливает трюк, начисляет score."],
        ["BallSpawner / BallPool / Ball", "Проп для трюков, использует ObjectPool и ReturnMe событие."],
    ],
    [4.2, 10.8],
)

h(doc, 1, "11. Object Pool и окружение")
doc.add_paragraph(
    "ObjectPool<T> используется для переиспользования объектов, например мячей/частиц. "
    "Environment-компоненты сейчас простые: JumpPad добавляет force rigidbody, InsideHouseZone работает как trigger-зона."
)

h(doc, 1, "12. Как добавить новый груз")
numbered(doc, "Создайте новый класс, наследник Cargo, или используйте один из готовых типов в SpecialCargos.cs.")
numbered(doc, "Если нужна runtime-память груза, храните ее в ActiveCargo.EffectTimer, SecondaryTimer, State или расширьте ActiveCargo аккуратно.")
numbered(doc, "Для эффекта при подборе переопределите OnPickup.")
numbered(doc, "Для периодической логики переопределите OnTick.")
numbered(doc, "Для реакции на столкновения используйте OnPlayerCollision.")
numbered(doc, "Для изменения награды используйте CalculateValue или GetGlobalRewardMultiplier.")
numbered(doc, "Создайте asset через Create > Cargo > ... и назначьте его в CargoPickup.")
numbered(doc, "Если груз должен отображаться в UI, назначьте Icon в asset.")

h(doc, 1, "13. Практические зависимости и правила")
table(
    doc,
    ["Правило", "Почему важно"],
    [
        ["Не переносить уникальную логику груза в PlayerCargoModule.", "Модуль должен быть диспетчером, а поведение должно жить в Cargo-наследниках."],
        ["Не писать напрямую в MaxSpeed/gravity из грузов.", "Иначе один груз может стереть эффект другого. Используйте Add/Remove множители Player."],
        ["Не уничтожать DeliveryPoint из DeliveryPoint.", "Жизненный цикл общей точки контролирует CargoManager через PlayerCargoModule."],
        ["Для новых UI грузов подписываться на CargoAdded/CargoRemoved/ComboChanged.", "Так UI не зависит от внутреннего списка напрямую."],
        ["Для новых врагов вызывайте Player.NotifyCargoScoreDamage после урона.", "VIP/броня/хрупкие грузы должны знать о damage-событиях."],
        ["Для новых pickup объектов используйте CargoPickup.", "Стрелка ищет ближайшие активные CargoPickup через статический список."],
    ],
    [5.0, 9.8],
)

h(doc, 1, "14. Быстрый чеклист настройки сцены")
bullet(doc, "Player prefab: должен иметь Player, PlayerCargoModule, PlayerCollisionHandler, ArcadeVehicleController, UiPlayer, ScoreUI.")
bullet(doc, "Player: назначить CargoManager и CargoArrowUI в inspector.")
bullet(doc, "CargoManager: назначить deliverySpawnPoint и deliveryPointPrefab.")
bullet(doc, "DeliveryPoint prefab: Collider должен быть trigger.")
bullet(doc, "CargoPickup prefab: Collider должен быть trigger, cargoData должен указывать на Cargo asset.")
bullet(doc, "CargoInventoryUI: назначить iconsRoot, cargoIconPrefab, comboText.")
bullet(doc, "ContrabandCargo: назначить policePrefab.")
bullet(doc, "Police prefab: для ElectricCargo желательно иметь tag Police или AICarChase на объекте/детях.")

doc.add_page_break()
h(doc, 1, "15. Известные риски и места для будущей чистки")
table(
    doc,
    ["Место", "Риск", "Рекомендация"],
    [
        ["SpecialCargos.cs", "Много классов в одном файле удобно для старта, но файл станет большим.", "Позже разнести каждый груз в отдельный файл."],
        ["CargoFreezeEffect", "Прямое isKinematic может конфликтовать с AI/физикой, если есть другая система заморозки.", "Объединить с TimeStopManager или сделать общий status-effect слой."],
        ["EnemyCollisionHandler", "Урон считается как int от float _damage.", "Если нужен точный баланс, перейти на int или отдельную damage-модель."],
        ["CarGameLoader", "Вероятная ошибка присваивания _dontTryLoad = true.", "Исправить на `if (_dontTryLoad == true)` или `if (_dontTryLoad)`."],
        ["Cargo assets", "Новые классы есть, но assets/prefabs нужно создать и настроить в Unity.", "Создать набор ScriptableObject assets в Assets/5_CargoData."],
    ],
    [4.0, 5.8, 5.0],
)

doc.add_page_break()
h(doc, 1, "Приложение A. Основные файлы")
table(
    doc,
    ["Файл", "Зачем открывать"],
    [
        ["Assets/2_script/Player/Player.cs", "Главная точка связи игрока, очков, бустов и грузов."],
        ["Assets/2_script/Cargo/PlayerCargoModule.cs", "Главная логика активных грузов и комбо."],
        ["Assets/2_script/Cargo/Cargo.cs", "Контракт для всех типов грузов."],
        ["Assets/2_script/Cargo/SpecialCargos.cs", "Все новые типы грузов."],
        ["Assets/2_script/Cargo/CargoInventoryUI.cs", "UI списка грузов и счетчика комбо."],
        ["Assets/2_script/Cargo/CargoArrowUI.cs", "Стрелка на ближайший груз."],
        ["Assets/2_script/BoostsLogic/BoostSystem.cs", "Менеджер runtime-бустов."],
        ["Assets/2_script/Player/PlayerUI/ScoreSystem.cs", "Система очков."],
        ["Assets/2_script/Enemy]/AICarChase.cs", "AI полицейской машины."],
        ["Assets/2_script/CarLoader/GarageUIManager.cs", "UI выбора машины."],
    ],
    [7.0, 7.8],
)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("CrazyCity project documentation")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(120, 120, 120)

doc.save(OUT)
print(OUT)
